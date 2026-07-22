from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from pydantic import BaseModel
from typing import Dict, Any, List
import os
import shutil
from sqlalchemy.orm import Session

from database.core import get_db
from database.models import Project, Document, ExtractedField
from services.document_processing import pdf_to_images, preprocess_image
from services.ocr_service import extract_text_from_image
from services.classifier import classify_document

# =========================================================
# MULTI-FORMAT EXTRACTION UTILITIES
# =========================================================

def clean_extracted_text(text):
    if text is None:
        return ""
    text = str(text)
    text = text.replace("\x00", " ")
    lines = [line.rstrip() for line in text.splitlines()]
    return "\n".join(lines).strip()

def extract_from_docx(file_path):
    from docx import Document
    document = Document(file_path)
    extracted_text = []
    for para in document.paragraphs:
        if para.text.strip():
            extracted_text.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(clean_extracted_text(cell.text))
            extracted_text.append(" | ".join(row_text))
    return "\n".join(extracted_text)

def extract_from_excel(file_path):
    import pandas as pd
    extracted_text = []
    extension = os.path.splitext(file_path)[1].lower()
    if extension == ".csv":
        try:
            df = pd.read_csv(file_path, dtype=str, encoding="utf-8", encoding_errors="ignore")
        except Exception:
            df = pd.read_csv(file_path, dtype=str, encoding="latin-1", encoding_errors="ignore")
        extracted_text.append(df.fillna("").to_string(index=False))
    else:
        excel_file = pd.ExcelFile(file_path)
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, dtype=str)
            df = df.fillna("")
            if not df.empty:
                extracted_text.append(df.to_string(index=False))
    return "\n".join(extracted_text)

def extract_from_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
            return f.read()

def extract_from_eml(file_path):
    from email import policy
    from email.parser import BytesParser
    from bs4 import BeautifulSoup
    extracted_text = []
    with open(file_path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)
    extracted_text.append(f"Subject: {msg.get('subject', '')}")
    extracted_text.append(f"From: {msg.get('from', '')}")
    extracted_text.append(f"To: {msg.get('to', '')}")
    extracted_text.append(f"Cc: {msg.get('cc', '')}")
    extracted_text.append(f"Date: {msg.get('date', '')}")
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            try:
                if content_type == "text/plain":
                    extracted_text.append(part.get_content())
                elif content_type == "text/html":
                    html = part.get_content()
                    soup = BeautifulSoup(html, "html.parser")
                    extracted_text.append(soup.get_text(separator="\n"))
            except Exception:
                continue
    else:
        content_type = msg.get_content_type()
        if content_type == "text/plain":
            extracted_text.append(msg.get_content())
        elif content_type == "text/html":
            html = msg.get_content()
            soup = BeautifulSoup(html, "html.parser")
            extracted_text.append(soup.get_text(separator="\n"))
    return "\n".join(extracted_text)

def extract_from_msg(file_path):
    import extract_msg
    msg = extract_msg.Message(str(file_path))
    extracted_text = [
        f"Subject: {msg.subject or ''}",
        f"From: {msg.sender or ''}",
        f"To: {msg.to or ''}",
        f"Cc: {msg.cc or ''}",
        f"Date: {msg.date or ''}",
        msg.body or ""
    ]
    return "\n".join(extracted_text)

TEXT_EXTENSIONS = {".docx", ".xlsx", ".xls", ".csv", ".txt", ".eml", ".msg"}

router = APIRouter(prefix="/processing", tags=["Processing"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

CURRENT_PROJECT_ID = None

@router.post("/upload_and_classify")
async def upload_and_classify(
    project_id: int = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Uploads a PDF/Image, runs OCR, classifies, extracts structured fields, and saves to DB.
    """
    global CURRENT_PROJECT_ID
    if CURRENT_PROJECT_ID is None or CURRENT_PROJECT_ID != project_id:
        print(f"New upload session detected (project_id={project_id}). Clearing old temporary case JSON files...")
        DATA_STORE_DIR = os.getenv("DATA_STORE_DIR", os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data_store"))
        if os.path.exists(DATA_STORE_DIR):
            for filename in os.listdir(DATA_STORE_DIR):
                if filename.endswith(".json"):
                    try:
                        os.remove(os.path.join(DATA_STORE_DIR, filename))
                    except Exception as e:
                        print(f"Error removing case file {filename}: {e}")
        CURRENT_PROJECT_ID = project_id

    file_ext = os.path.splitext(file.filename)[1].lower()
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    all_blocks = []
    use_ocr = True
    image_paths = []

    if file_ext in TEXT_EXTENSIONS:
        try:
            extracted_text = ""
            if file_ext == ".docx":
                extracted_text = extract_from_docx(file_path)
            elif file_ext in [".xlsx", ".xls", ".csv"]:
                extracted_text = extract_from_excel(file_path)
            elif file_ext == ".txt":
                extracted_text = extract_from_txt(file_path)
            elif file_ext == ".eml":
                extracted_text = extract_from_eml(file_path)
            elif file_ext == ".msg":
                extracted_text = extract_from_msg(file_path)
            
            extracted_text = clean_extracted_text(extracted_text)
            lines = extracted_text.splitlines()
            for line_idx, line in enumerate(lines):
                line_str = line.strip()
                if line_str:
                    all_blocks.append({
                        "text": line_str,
                        "bbox": [0, 0, 100, 100],  # Default bounding box
                        "page": 1,
                        "page_width": 100,
                        "page_height": 100
                    })
            use_ocr = False
        except Exception as e:
            print(f"Error extracting text from non-image file {file.filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to process text document: {str(e)}")
    elif file_ext == ".pdf":
        try:
            image_paths = pdf_to_images(file_path, UPLOAD_DIR)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to process PDF: {str(e)}")
    else:
        image_paths = [file_path]

    # Try fast native text extraction if PDF
    if file_ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(file_path)
            temp_blocks = []
            for page_idx in range(len(doc)):
                page = doc.load_page(page_idx)
                rect = page.rect
                page_width = rect.width
                page_height = rect.height
                blocks = page.get_text("blocks")
                for b in blocks:
                    text = b[4].strip()
                    if text:
                        temp_blocks.append({
                            "text": text,
                            "bbox": [int(b[0]), int(b[1]), int(b[2]), int(b[3])],
                            "page": page_idx + 1,
                            "page_width": page_width,
                            "page_height": page_height
                        })
            # Verify if significant text was extracted
            full_temp_text = " ".join([b["text"] for b in temp_blocks])
            if len(full_temp_text.strip()) > 100:
                all_blocks = temp_blocks
                use_ocr = False
        except Exception as e:
            print(f"Error extracting native PDF text: {e}")

    # Fallback to PaddleOCR if not a PDF or if native text was insufficient (scanned)
    if use_ocr:
        for page_idx, img_path in enumerate(image_paths):
            processed_path = preprocess_image(img_path)
            blocks = extract_text_from_image(processed_path)
            from PIL import Image
            img_width, img_height = 1, 1
            if os.path.exists(processed_path):
                with Image.open(processed_path) as img:
                    img_width, img_height = img.size
            for b in blocks:
                b["page"] = page_idx + 1
                b["page_width"] = img_width
                b["page_height"] = img_height
            all_blocks.extend(blocks)

    full_text = "\n".join([b["text"] for b in all_blocks])
    classification = classify_document(full_text, filename=file.filename)

    # Fallback to PaddleOCR if native PDF text was garbled or matched no categories
    if classification == "Unknown" and not use_ocr and file_ext == ".pdf":
        print(f"Native text classification returned 'Unknown' for {file.filename}. Falling back to OCR...")
        use_ocr = True
        all_blocks = []
        for page_idx, img_path in enumerate(image_paths):
            processed_path = preprocess_image(img_path)
            blocks = extract_text_from_image(processed_path)
            from PIL import Image
            img_width, img_height = 1, 1
            if os.path.exists(processed_path):
                with Image.open(processed_path) as img:
                    img_width, img_height = img.size
            for b in blocks:
                b["page"] = page_idx + 1
                b["page_width"] = img_width
                b["page_height"] = img_height
            all_blocks.extend(blocks)
        full_text = "\n".join([b["text"] for b in all_blocks])
        classification = classify_document(full_text, filename=file.filename)
    
    # 1. AI Classification of Subcategory based on CSV rule catalogue
    from services.ai_extractor import classify_subcategory_ai, extract_entities_ai
    subcategory = classify_subcategory_ai(full_text)
    
    # 2. AI Entity Extraction
    extracted_fields = extract_entities_ai(classification, all_blocks)
    
    # Normalize absolute coordinates to integer percentages
    for field_name, field_data in extracted_fields.items():
        if isinstance(field_data, dict) and "bbox" in field_data:
            bbox = field_data["bbox"]
            width = field_data.get("page_width") or 1
            height = field_data.get("page_height") or 1
            x1 = int(round(max(0.0, min(100.0, (bbox[0] / width) * 100.0))))
            y1 = int(round(max(0.0, min(100.0, (bbox[1] / height) * 100.0))))
            x2 = int(round(max(0.0, min(100.0, (bbox[2] / width) * 100.0))))
            y2 = int(round(max(0.0, min(100.0, (bbox[3] / height) * 100.0))))
            field_data["bbox"] = [x1, y1, x2, y2]

    extracted_fields["raw_text"] = {"field": "raw_text", "value": full_text, "page": 1, "bbox": [0, 0, 0, 0]}
    
    # 3. Correlate and Store in respective JSON files
    from services.correlation import correlate_and_store_document, DuplicateDocumentError
    try:
        case_data = correlate_and_store_document(classification, extracted_fields, subcategory)
    except DuplicateDocumentError as e:
        raise HTTPException(status_code=400, detail=str(e))
        
    # Create Document record
    db_doc = Document(project_id=project_id, document_type=classification, file_name=file.filename, file_path=file_path)
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    
    # Save to ExtractedField DB for Phase 16 Evidence Tracking
    for field_name, field_data in extracted_fields.items():
        if field_data and field_data.get("value") is not None:
            bbox = field_data.get("bbox", [0, 0, 0, 0])
            ext_field = ExtractedField(
                project_id=project_id,
                document_id=db_doc.id,
                field_name=field_name,
                field_value=str(field_data["value"]),
                page=field_data.get("page", 1),
                bbox_x1=bbox[0],
                bbox_y1=bbox[1],
                bbox_x2=bbox[2],
                bbox_y2=bbox[3]
            )
            db.add(ext_field)
    
    db.commit()
    
    image_url = f"/uploads/{os.path.basename(image_paths[0])}" if image_paths else None
    image_urls = [f"/uploads/{os.path.basename(p)}" for p in image_paths] if image_paths else []

    # Merge overall case data into response extracted fields so UI gets the full correlated context
    merged_fields = {}
    for doc_t, doc_fields in case_data.get("documents", {}).items():
        for f_name, f_val in doc_fields.items():
            if f_name != "raw_text":
                merged_fields[f_name] = f_val

    return {
        "document_id": db_doc.id,
        "filename": file.filename,
        "classification": classification,
        "extracted_fields": extracted_fields,
        "merged_fields": merged_fields,
        "image_url": image_url,
        "image_urls": image_urls,
        "raw_text": full_text
    }

class RulesEngineRequest(BaseModel):
    project_id: int
    extracted_data: Dict[str, Dict[str, Any]]

@router.post("/run_rules")
async def run_rules(request: RulesEngineRequest, db: Session = Depends(get_db)):
    """
    Executes validation rules against the structured payload, including Duplicate Detection and Approval Validation.
    """
    from validators.engine import run_validation_engine_with_meta
    res_meta = run_validation_engine_with_meta(request.extracted_data, db=db, current_project_id=request.project_id)
    
    return {
        "overall_status": res_meta["overall_status"],
        "rule_results": res_meta["rule_results"],
        "subcategory": res_meta["subcategory"]
    }

